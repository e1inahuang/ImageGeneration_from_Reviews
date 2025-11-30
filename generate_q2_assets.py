
import pandas as pd
import numpy as np
import json
import matplotlib.pyplot as plt
from sentence_transformers import SentenceTransformer
from sklearn.cluster import KMeans
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Generating Q2 Assets...")

# 1. Load Data
df = pd.read_csv('cleaned_helmet_reviews.csv')
reviews = df['review'].tolist()

# 2. Generate Embeddings
print("Generating embeddings...")
model = SentenceTransformer('all-MiniLM-L6-v2')
embeddings = model.encode(reviews, show_progress_bar=True)

# 3. Clustering
print("Clustering...")
n_clusters = 8
kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
cluster_labels = kmeans.fit_predict(embeddings)
df['cluster'] = cluster_labels

# 4. Generate Plot
print("Generating plot...")
plt.figure(figsize=(10, 6))
cluster_counts = df['cluster'].value_counts().sort_index()
bars = plt.bar(cluster_counts.index, cluster_counts.values, color='#4A90E2', alpha=0.9)
plt.xlabel('Cluster ID', fontsize=12)
plt.ylabel('Number of Reviews', fontsize=12)
plt.title('Review Distribution Across Topics (Clusters)', fontsize=14, fontweight='bold')
plt.grid(axis='y', alpha=0.3)
plt.xticks(range(n_clusters))

# Add counts on top
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 2,
             f'{int(height)}',
             ha='center', va='bottom')

plt.tight_layout()
plt.savefig('cluster_distribution.png', dpi=300)
print("Saved cluster_distribution.png")

# 5. Load Extracted Features
with open('helmet_extracted_features.json', 'r') as f:
    features = json.load(f)

# 6. Create DOCX
print("Creating DOCX...")
doc = Document()

# Title
heading = doc.add_heading('Q2: GenAI Analysis of Customer Reviews', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Methodology', level=1)
p = doc.add_paragraph(
    "In this section, we applied Generative AI techniques to analyze the 582 customer reviews "
    "for the VICTGOAL Bike Helmet. The workflow involved:"
)
ul = doc.add_paragraph()
ul.style = 'List Bullet'
ul.add_run("Generating sentence embeddings for all reviews using the 'all-MiniLM-L6-v2' model.")
doc.add_paragraph("Performing K-means clustering to identify distinct review topics.", style='List Bullet')
doc.add_paragraph("Using FAISS (Facebook AI Similarity Search) to retrieve semantically relevant reviews.", style='List Bullet')
doc.add_paragraph("Leveraging Large Language Models (LLM) to extract detailed visual and functional features.", style='List Bullet')
doc.add_paragraph("Generating a product visualization based on the extracted insights.", style='List Bullet')

# Cluster Analysis
doc.add_heading('2. Topic Clustering Analysis', level=1)
doc.add_paragraph(
    "We identified 8 distinct clusters of reviews, representing different topics and sentiments. "
    "The distribution of reviews across these clusters is shown below:"
)

# Add Plot
doc.add_picture('cluster_distribution.png', width=Inches(6.0))
last_paragraph = doc.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Cluster Descriptions (Based on our manual analysis of samples)
doc.add_heading('Cluster Insights:', level=2)
cluster_descriptions = [
    "General neutral/short feedback.",
    "Positive feedback on fit and delivery speed.",
    "Focus on the rear LED light and safety features.",
    "Mixed reviews regarding visor assembly and fitment.",
    "High praise for value for money and comfort.",
    "Specific complaints about assembly difficulties.",
    "Detailed praise for style, goggles, and versatility.",
    "General enthusiasm and recommendations."
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Cluster ID'
hdr_cells[1].text = 'Dominant Theme'

for i, desc in enumerate(cluster_descriptions):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = desc

# Feature Extraction
doc.add_heading('3. GenAI Feature Extraction', level=1)
doc.add_paragraph(
    "Using an LLM to analyze the most relevant reviews, we extracted the following key product attributes:"
)

# Design
doc.add_heading('Design & Aesthetics', level=2)
doc.add_paragraph(f"Style: {features['design']['overall_style']}")
doc.add_paragraph(f"Shape: {features['design']['shape']}")

# Features
doc.add_heading('Key Features', level=2)
doc.add_paragraph(f"Magnetic Goggles: {features['key_features']['goggles']}")
doc.add_paragraph(f"Rear Light: {features['key_features']['light']}")
doc.add_paragraph(f"Ventilation: {features['key_features']['ventilation']}")

# Materials
doc.add_heading('Materials & Build', level=2)
doc.add_paragraph(f"Shell/Foam: {features['materials']['shell']} / {features['materials']['foam']}")
doc.add_paragraph(f"Padding: {features['materials']['padding']}")

# Image Generation
doc.add_heading('4. AI-Generated Product Visualization', level=1)
doc.add_paragraph(
    "Based on the extracted features, the following image prompt was generated and used to create a visualization:"
)

p_prompt = doc.add_paragraph()
p_prompt.add_run("Prompt: ").bold = True
p_prompt.add_run(features['image_prompt']).italic = True

doc.add_paragraph(
    "[Note: Image generation was attempted using DALL-E 3. Please insert the generated image here if available, "
    "or use the prompt above to generate one.]"
)

# Save
doc.save('Final_Project_Q2.docx')
print("Saved Final_Project_Q2.docx")
