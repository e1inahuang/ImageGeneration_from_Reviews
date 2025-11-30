"""
Q1 Documentation Generator for 94844 Final Project
Generates DOCX file with comprehensive Q1 analysis of the VICTGOAL Bike Helmet
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('94844 Final Project - Product Review Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Group Assignment - Fall 2025')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.color.rgb = RGBColor(64, 64, 64)

doc.add_paragraph()

# Add group information
group_info = doc.add_paragraph()
group_info.add_run('Product Selected: ').bold = True
group_info.add_run('VICTGOAL Bike Helmet with Visor and Goggles\n')
group_info.add_run('Other Group Products: ').bold = True
group_info.add_run('Egg Cooker, Cat Tree House')

doc.add_page_break()

# ==============================
# Q1: PRODUCT AND DATA UNDERSTANDING
# ==============================

doc.add_heading('Q1: Product and Data Understanding', 1)

doc.add_paragraph(
    'This section provides a comprehensive overview of the VICTGOAL Bike Helmet, including detailed '
    'product specifications, target audience analysis, and an initial exploration of the customer review data.'
)

# Product Description
doc.add_heading('1. Product Description', 2)

product_desc = doc.add_paragraph()
product_desc.add_run('Product Name: ').bold = True
product_desc.add_run('VICTGOAL Bike Helmet with Visor and Goggles\n')
product_desc.add_run('Product URL: ').bold = True
product_desc.add_run('https://www.walmart.com/ip/VICTGOAL-Bike-Helmet-with-Visor-and-Goggles-USB-Rechargeable-Light-for-Adult-Men-Women/732290255\n\n')

doc.add_heading('Key Features:', 3)

features = [
    'Safety Certification: Complies with U.S. CPSC Safety Standard for Bicycle Helmets',
    'Construction: High-density EPS foam structure with PC shell for superior impact absorption',
    'Eye Protection: Detachable visor and magnetic goggles that block glare, sand, bugs, and wind',
    'Visibility: Super bright USB rechargeable rear light with 3 flashing modes (10-hour battery life after 2-hour charge)',
    'Ventilation: 21 breathable vents to reduce air resistance and maintain cooling',
    'Weight: Ultra-lightweight at just 310g (0.68 lb)',
    'Adjustability: One-hand adjustable sizing for head circumference 57-61 cm (22.4-24 inches)',
    'Comfort: Breathable inner padding with removable, washable liner'
]

for feature in features:
    p = doc.add_paragraph(feature, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_heading('Target Audience:', 3)

target_audience = doc.add_paragraph(
    'The VICTGOAL Bike Helmet is designed for adult men and women who engage in various cycling activities, '
    'including road biking, mountain biking, and urban commuting. The product particularly appeals to safety-conscious '
    'cyclists who value additional features like integrated lighting, eye protection, and comfortable fit. The helmet '
    'is suitable for both recreational riders and those who commute regularly by bicycle.'
)

doc.add_heading('Product Specifications:', 3)

specs_table = doc.add_table(rows=5, cols=2)
specs_table.style = 'Light Grid Accent 1'

specs_data = [
    ('Weight', '0.68 lb (310g)'),
    ('Material', 'EPS foam, PC shell'),
    ('Bicycle Type', 'Road Bikes, Mountain Bikes'),
    ('Size Range', '57-61 cm (22.4-24 inches)'),
    ('Warranty', '90-day limited warranty against manufacturing defects')
]

for i, (spec, value) in enumerate(specs_data):
    row = specs_table.rows[i]
    row.cells[0].text = spec
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[1].text = value

doc.add_paragraph()

# Data Description
doc.add_heading('2. Data Description', 2)

data_desc = doc.add_paragraph()
data_desc_text = (
    'The review dataset for the VICTGOAL Bike Helmet was collected from multiple sources, including Walmart.com '
    'and the manufacturer\'s website (victgoal.com). The data was compiled in CSV format with reviews spanning '
    'from 2019 to 2025.\n\n'
)
data_desc.add_run(data_desc_text)

doc.add_heading('Dataset Structure:', 3)

dataset_info = doc.add_paragraph()
dataset_info.add_run('Source File: ').bold = True
dataset_info.add_run('reviews.csv\n')
dataset_info.add_run('Review Column: ').bold = True
dataset_info.add_run('"tl-m" (4th column)\n')
dataset_info.add_run('Rating Column: ').bold = True
dataset_info.add_run('"w_iUH7" (3rd column)\n')
dataset_info.add_run('Date Column: ').bold = True
dataset_info.add_run('"f7" (1st column)\n')
dataset_info.add_run('Reviewer Column: ').bold = True
dataset_info.add_run('"f7 2" (2nd column)\n\n')

dataset_info.add_run('Review Sources: ').bold = True
dataset_info.add_run('Reviews were collected from both Walmart.com and victgoal.com, providing a comprehensive '
                     'view of customer feedback across different platforms.')

doc.add_page_break()

# Initial Data Exploration
doc.add_heading('3. Initial Data Exploration', 2)

exploration_intro = doc.add_paragraph(
    'A preliminary analysis of the review dataset reveals important patterns about customer satisfaction '
    'and engagement with the VICTGOAL Bike Helmet. Below are the key statistics:'
)

doc.add_heading('Overall Statistics:', 3)

stats_para = doc.add_paragraph()
stats_para.add_run('Total Number of Reviews: ').bold = True
stats_para.add_run('582\n')
stats_para.add_run('Review Period: ').bold = True
stats_para.add_run('2019 - 2025 (approximately 6 years)\n')
stats_para.add_run('Average Reviews per Year: ').bold = True
stats_para.add_run('~97 reviews/year\n\n')

doc.add_heading('Rating Distribution:', 3)

# Create rating distribution table
rating_table = doc.add_table(rows=6, cols=3)
rating_table.style = 'Medium Grid 1 Accent 1'

# Header row
header_cells = rating_table.rows[0].cells
header_cells[0].text = 'Rating'
header_cells[1].text = 'Count'
header_cells[2].text = 'Percentage'

for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
rating_data = [
    ('5 stars', 469, '80.6%'),
    ('4 stars', 62, '10.7%'),
    ('3 stars', 25, '4.3%'),
    ('2 stars', 13, '2.2%'),
    ('1 star', 13, '2.2%')
]

for i, (rating, count, percentage) in enumerate(rating_data, 1):
    row = rating_table.rows[i]
    row.cells[0].text = rating
    row.cells[1].text = str(count)
    row.cells[2].text = percentage
    
    # Center align count and percentage
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Rating insights
rating_insights = doc.add_paragraph()
rating_insights.add_run('Key Insights:\n').bold = True
insights = [
    'Overwhelming Positive Reception: 80.6% of customers gave the helmet 5 stars, indicating high satisfaction',
    'Strong Overall Rating: 91.3% of reviews are 4 stars or higher',
    'Low Negative Feedback: Only 4.4% of reviews are 1-2 stars, suggesting few serious issues',
    'Consistent Quality: The high percentage of top ratings suggests reliable product performance'
]

for insight in insights:
    p = doc.add_paragraph(insight, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_paragraph()

# Review Length Statistics
doc.add_heading('Review Length Analysis:', 3)

length_para = doc.add_paragraph()
length_para.add_run('Minimum Review Length: ').bold = True
length_para.add_run('1 word\n')
length_para.add_run('Maximum Review Length: ').bold = True
length_para.add_run('360 words\n')
length_para.add_run('Average Review Length: ').bold = True
length_para.add_run('36.0 words\n')
length_para.add_run('Median Review Length: ').bold = True
length_para.add_run('20.0 words\n\n')

length_insights = doc.add_paragraph()
length_insights.add_run('Analysis:\n').bold = True
length_insights.add_run(
    'The review length statistics reveal that while most customers provide concise feedback (median of 20 words), '
    'there is a significant range in detail level. The average of 36 words suggests that customers generally '
    'provide substantive feedback beyond simple ratings. The presence of reviews up to 360 words indicates that '
    'some customers are highly engaged and willing to share detailed experiences. This variety in review length '
    'will be valuable for different types of analysis, from sentiment analysis on brief reviews to detailed '
    'feature extraction from longer, more comprehensive reviews.'
)

doc.add_paragraph()

# Common Themes (preliminary)
doc.add_heading('Preliminary Theme Identification:', 3)

themes_intro = doc.add_paragraph(
    'Based on initial review scanning, several recurring themes emerge from customer feedback:'
)

themes = [
    ('Magnetic Visor/Goggles', 'Customers frequently mention and appreciate the magnetic attachment system for the visor, '
     'noting its convenience and functionality for eye protection'),
    
    ('Rear Safety Light', 'The USB rechargeable rear light is highlighted as a valuable safety feature, particularly '
     'for visibility during low-light conditions'),
    
    ('Comfort and Fit', 'Many reviews emphasize the helmet\'s comfortable fit, lightweight design, and easy '
     'adjustability using the dial system'),
    
    ('Ventilation', 'The 21 vents are praised for keeping the head cool during rides, especially in warm weather'),
    
    ('Value for Money', 'Reviewers frequently compare the helmet favorably to more expensive alternatives, noting '
     'the combination of features at an affordable price point'),
    
    ('Crash Protection', 'Several reviews mention the helmet performing well in actual crashes, protecting the '
     'wearer\'s head as intended')
]

for theme, description in themes:
    p = doc.add_paragraph()
    p.add_run(f'{theme}: ').bold = True
    p.add_run(description)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

# Data Quality Notes
doc.add_heading('Data Quality Observations:', 3)

quality_notes = [
    'Complete Data: All 582 reviews contain both rating and text components',
    'Time Span: Reviews span approximately 6 years, providing longitudinal perspective on product performance',
    'Multi-Source: Data from both retail (Walmart) and manufacturer (victgoal.com) platforms offers diverse perspectives',
    'Authenticity Indicators: Reviews show varied writing styles, detailed personal experiences, and both positive and negative feedback, suggesting authentic customer voices',
    'Review Verification: Some reviews indicate verified purchases and actual product usage (mention of crashes, long-term use, etc.)'
]

for note in quality_notes:
    p = doc.add_paragraph(note, style='List Bullet')
    p.paragraph_format.left_indent = Inches (0.5)

doc.add_paragraph()

#  Summary
doc.add_heading('Q1 Summary', 2)

summary_text = (
    'The VICTGOAL Bike Helmet represents a well-featured safety product targeting adult cyclists across various '
    'cycling disciplines. The review dataset of 582 customer opinions reveals an overwhelmingly positive reception, '
    'with over 80% of customers awarding the maximum 5-star rating. The helmet\'s key differentiating features—'
    'magnetic goggles/visor, USB rechargeable rear light, and lightweight yet protective construction—are '
    'consistently mentioned as value drivers.\n\n'
    
    'The data quality is strong, with complete reviews spanning multiple years and platforms. The variety in '
    'review length (1-360 words) provides both quick sentiment snapshots and detailed usage experiences. '
    'Preliminary theme analysis suggests that customers particularly appreciate the innovative design features '
    '(magnetic visor system), safety components (rear light), comfort factors (fit, weight, ventilation), and '
    'overall value proposition.\n\n'
    
    'This comprehensive dataset provides an excellent foundation for deeper generative AI analysis in Q2, where '
    'we can extract more nuanced insights about specific features, use cases, and customer segments.'
)

summary_para = doc.add_paragraph(summary_text)

# Save the document
doc.save('Final_Project_Q1.docx')

print("✓ Q1 documentation successfully created: Final_Project_Q1.docx")
print("  - Total pages: ~4-5 pages")
print("  - Sections: Product Description, Data Description, Initial Exploration")
print("  - Statistics: 582 reviews analyzed")
print("  - Rating distribution, review length, and preliminary themes documented")
