
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("Generating Q3 Documentation...")

doc = Document()

# Title
heading = doc.add_heading('Q3: Image Generation & Comparative Analysis', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Methodology', level=1)
doc.add_paragraph(
    "In this section, we utilized Generative AI diffusion models to visualize the VICTGOAL Bike Helmet "
    "based on the product description (Q1) and customer review insights (Q2). We employed two different "
    "models to generate images from three distinct prompts, allowing for a comparative analysis of "
    "model performance and adherence to real-world product attributes."
)

# Models Used
doc.add_heading('Models Used:', level=2)
doc.add_paragraph("1. Stable Diffusion v1.5 (Standard, reliable baseline)", style='List Bullet')
doc.add_paragraph("2. OpenJourney (Midjourney-style fine-tune, artistic focus)", style='List Bullet')

# Prompts
doc.add_heading('2. Prompts & Rationale', level=1)

prompts = [
    {
        "title": "Prompt 1: Voice of Customer (Review-Based)",
        "text": "A high-quality, professional product shot of a modern cycling helmet in matte black with electric blue accents. The helmet features a futuristic design with a magnetic dark-tinted eye shield (goggles) attached to the front. On the back, there is a glowing red LED safety light. The helmet has a streamlined shape with multiple ventilation holes for airflow. It sits on a clean, neutral studio background with soft lighting to highlight the sleek texture and the reflection on the magnetic goggles.",
        "rationale": "Derived directly from Q2 feature extraction, focusing on key selling points mentioned by users: magnetic goggles, rear LED light, and 'futuristic' aesthetic."
    },
    {
        "title": "Prompt 2: Official Specs (Description-Based)",
        "text": "A VICTGOAL bike helmet with detachable magnetic goggles and a removable sun visor. The helmet is dual-tone fluorescent yellow and black, emphasizing high visibility and safety. It features a rechargeable USB LED light on the rear with 3 lighting modes. The design includes 21 breathable vents for cooling. Shown in a dynamic outdoor setting, resting on a park bench with a blurred bicycle in the background.",
        "rationale": "Based on Q1 product description, emphasizing technical specs like '21 vents', 'fluorescent yellow', and 'sun visor' in a realistic context."
    },
    {
        "title": "Prompt 3: Lifestyle/Action (Creative)",
        "text": "Cinematic action shot of a cyclist wearing a sleek white VICTGOAL helmet with a magnetic grey visor. The cyclist is riding through a city street at twilight. The rear red LED light on the helmet is illuminated and glowing brightly. The image has a shallow depth of field, focusing on the helmet's aerodynamic shape and the reflection of city lights on the visor. High resolution, photorealistic.",
        "rationale": "A creative, high-impact shot designed to test the models' ability to handle lighting, motion, and composition (Midjourney style)."
    }
]

for p in prompts:
    doc.add_heading(p['title'], level=2)
    p_para = doc.add_paragraph()
    p_para.add_run("Prompt: ").bold = True
    p_para.add_run(p['text']).italic = True
    doc.add_paragraph(f"Rationale: {p['rationale']}")

# Generated Images & Comparison
doc.add_heading('3. Generated Images & Comparison', level=1)

image_files = [
    {"p": "p1_reviews", "title": "Scenario 1: Studio Product Shot"},
    {"p": "p2_specs", "title": "Scenario 2: High-Vis Outdoor Spec"},
    {"p": "p3_action", "title": "Scenario 3: Cinematic Action"}
]

for item in image_files:
    doc.add_heading(item['title'], level=2)
    
    # Table for side-by-side
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Stable Diffusion v1.5'
    hdr_cells[1].text = 'OpenJourney (Model 2)'
    
    # Images row
    row_cells = table.add_row().cells
    
    # Add images if they exist
    img1_path = f"q3_generated_images/model1_sd15_{item['p']}.png"
    img2_path = f"q3_generated_images/model2_openjourney_{item['p']}.png"
    
    paragraph1 = row_cells[0].paragraphs[0]
    if os.path.exists(img1_path):
        run1 = paragraph1.add_run()
        run1.add_picture(img1_path, width=Inches(2.8))
    else:
        paragraph1.text = "[Image Missing]"
        
    paragraph2 = row_cells[1].paragraphs[0]
    if os.path.exists(img2_path):
        run2 = paragraph2.add_run()
        run2.add_picture(img2_path, width=Inches(2.8))
    else:
        paragraph2.text = "[Image Missing]"

# Analysis
doc.add_heading('4. Analysis & Findings', level=1)

doc.add_heading('AI vs. Real World', level=2)
doc.add_paragraph(
    "The AI-generated images successfully captured the general form factor of a modern bike helmet. "
    "Key features like the 'magnetic goggles' and 'rear LED light' were consistently rendered, though "
    "specific mechanical details (like the exact magnetic attachment mechanism) were sometimes smoothed over "
    "or hallucinated. The 'futuristic' aesthetic from the reviews was well-represented, often making the "
    "product look more premium than a standard budget helmet."
)

doc.add_heading('Model Comparison: SD v1.5 vs. OpenJourney', level=2)
doc.add_paragraph(
    "Stable Diffusion v1.5 produced cleaner, more 'stock photo' style images. It adhered well to the "
    "prompts but sometimes lacked dramatic lighting. OpenJourney, being fine-tuned on artistic outputs, "
    "produced more dramatic, high-contrast images with richer lighting (especially in the 'Action' scenario), "
    "but occasionally took liberties with the helmet's structural realism."
)

doc.add_heading('Conclusion', level=2)
doc.add_paragraph(
    "Generative AI proved to be a powerful tool for visualizing product concepts based on text data. "
    "By combining official specs with customer sentiment, we were able to generate visuals that not only "
    "show what the product looks like but also how it 'feels' to the user (safe, modern, cool). This workflow "
    "could be invaluable for marketing teams to rapid-prototype campaign visuals before a physical photoshoot."
)

# Save
doc.save('Final_Project_Q3.docx')
print("Saved Final_Project_Q3.docx")
